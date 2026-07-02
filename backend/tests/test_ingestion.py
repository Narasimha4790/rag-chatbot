import os
import pytest
from unittest.mock import MagicMock, patch
from backend.app.services.ingestion import extract_text, parse_txt, parse_docx, parse_csv, parse_pdf

def test_parse_txt(tmp_path):
    txt_file = tmp_path / "test.txt"
    txt_content = "Hello, this is a test text file."
    txt_file.write_text(txt_content, encoding="utf-8")
    
    parsed = parse_txt(str(txt_file))
    assert parsed == txt_content

def test_parse_csv(tmp_path):
    csv_file = tmp_path / "test.csv"
    csv_content = "Name,Age,Role\nAlice,30,Engineer\nBob,25,Designer"
    csv_file.write_text(csv_content, encoding="utf-8")
    
    parsed = parse_csv(str(csv_file))
    expected = (
        "Row 1: Name: Alice, Age: 30, Role: Engineer\n"
        "Row 2: Name: Bob, Age: 25, Role: Designer"
    )
    assert parsed == expected

def test_parse_docx(tmp_path):
    docx_file = tmp_path / "test.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Paragraph 1 content.")
    doc.add_paragraph("Paragraph 2 content.")
    doc.save(str(docx_file))
    
    parsed = parse_docx(str(docx_file))
    assert "Paragraph 1 content." in parsed
    assert "Paragraph 2 content." in parsed

@patch("backend.app.services.ingestion.PdfReader")
def test_parse_pdf(mock_pdf_reader):
    mock_reader_inst = MagicMock()
    mock_page_1 = MagicMock()
    mock_page_1.extract_text.return_value = "PDF Page 1 text."
    mock_page_2 = MagicMock()
    mock_page_2.extract_text.return_value = "PDF Page 2 text."
    
    mock_reader_inst.pages = [mock_page_1, mock_page_2]
    mock_pdf_reader.return_value = mock_reader_inst
    
    parsed = parse_pdf("mock.pdf")
    assert parsed == "PDF Page 1 text.\nPDF Page 2 text."

def test_extract_text_routing(tmp_path):
    # Test routing to TXT
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("Hello text file", encoding="utf-8")
    assert extract_text(str(txt_file)) == "Hello text file"
    
    # Test file not found
    with pytest.raises(FileNotFoundError):
        extract_text("nonexistent.txt")
        
    # Test unsupported extension
    bad_file = tmp_path / "test.xyz"
    bad_file.write_text("dummy content", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_text(str(bad_file))
